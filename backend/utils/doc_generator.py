import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_resume(data, filename="My_Resume.docx"):
    """
    Generate a polished ATS-friendly resume in Word format and save it in generated_resumes/
    """

    # Ensure output folder exists
    output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "generated_resumes")
    os.makedirs(output_dir, exist_ok=True)

    # Create filename with student name + timestamp
    student_name = data.get("name", "Student").replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if not filename:
        filename = f"{student_name}_{timestamp}.docx"
    else:
        name, ext = os.path.splitext(filename)
        filename = f"{student_name}_{timestamp}{ext}"

    filepath = os.path.join(output_dir, filename)

    # Create document
    doc = Document()

    # ===== Header: Name =====
    name_heading = doc.add_heading(data.get("name", "Student"), level=0)
    name_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ===== Contact Info =====
    contact = data.get("contact", {})
    contact_line = " | ".join(
        filter(
            None,
            [
                contact.get("phone"),
                contact.get("email"),
                contact.get("linkedin"),
                contact.get("github"),
            ],
        )
    )
    if contact_line:
        para = doc.add_paragraph(contact_line)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ===== Career Objective / Summary =====
    if data.get("summary"):
        doc.add_heading("Career Objective", level=1)
        doc.add_paragraph(data["summary"])

    # ===== Education =====
    if data.get("education"):
        doc.add_heading("Education", level=1)
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.add_run(f"{edu['degree']}").bold = True
            p.add_run(f" - {edu['institution']} ({edu['year']})")

    # ===== Skills =====
    if data.get("skills"):
        doc.add_heading("Skills", level=1)
        for skill in data["skills"]:
            doc.add_paragraph(skill, style="List Bullet")

    # ===== Training / Internships =====
    if data.get("training"):
        doc.add_heading("Training / Internships", level=1)
        for t in data["training"]:
            doc.add_paragraph(
                f"{t['title']} ({t['from']} to {t['to']})", style="List Bullet"
            )

    # ===== Projects =====
    if data.get("projects"):
        doc.add_heading("Projects", level=1)
        for p in data["projects"]:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(p["name"]).bold = True
            para.add_run(f" - {p['description']}")

    # ===== Certifications =====
    if data.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for c in data["certifications"]:
            doc.add_paragraph(f"{c['title']} ({c['year']})", style="List Bullet")

    # ===== Achievements =====
    if data.get("achievements"):
        doc.add_heading("Achievements", level=1)
        for a in data["achievements"]:
            doc.add_paragraph(a, style="List Bullet")

    # ===== Extra-curricular / Leadership =====
    if data.get("extracurricular"):
        doc.add_heading("Extra-curricular / Leadership", level=1)
        for e in data["extracurricular"]:
            doc.add_paragraph(e, style="List Bullet")

    # Save
    doc.save(filepath)
    return filepath
